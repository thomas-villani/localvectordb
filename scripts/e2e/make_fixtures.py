"""Generate real document fixtures for the end-to-end test scripts.

Creates a set of genuine files (Markdown, HTML, PDF, DOCX, XLSX, Python
source) under scripts/e2e/fixtures/. Each document covers a distinct topic so
that vector-search assertions in the e2e scripts can check that semantically
related queries retrieve the right document.

Usage:
    ./.venv/Scripts/python.exe scripts/e2e/make_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

FIXTURES_DIR = Path(__file__).parent / "fixtures"

# ---------------------------------------------------------------------------
# Real document content, one distinct topic per document.
# ---------------------------------------------------------------------------

SPACE_MD = """\
# A Brief History of Space Exploration

## The Space Race

The modern era of space exploration began on 4 October 1957, when the Soviet
Union launched Sputnik 1, the first artificial satellite to orbit the Earth.
The 83 kg aluminium sphere transmitted a simple radio beep, but its political
impact was enormous: it triggered the space race between the Soviet Union and
the United States that would define the following two decades.

Yuri Gagarin became the first human in space on 12 April 1961, completing a
single orbit of the Earth aboard Vostok 1 in 108 minutes. The United States
responded with Project Mercury and, ultimately, the Apollo programme.

## The Apollo Programme

Apollo 11 landed the first humans on the Moon on 20 July 1969. Neil Armstrong
and Buzz Aldrin spent 21.5 hours on the lunar surface at Tranquility Base
while Michael Collins orbited above in the command module Columbia. Between
1969 and 1972, six Apollo missions landed twelve astronauts on the Moon and
returned 382 kg of lunar rock and soil to Earth.

## The Shuttle Era and the ISS

The Space Shuttle flew 135 missions between 1981 and 2011. Despite the
Challenger and Columbia tragedies, the shuttle fleet assembled the
International Space Station, deployed the Hubble Space Telescope, and made
spaceflight almost routine. The ISS has been continuously inhabited since
November 2000 and orbits the Earth at roughly 400 km altitude, travelling at
7.66 km/s.

## The Commercial Era

Private companies transformed spaceflight economics in the 2010s. Reusable
first-stage boosters cut launch costs dramatically, and commercial crew
vehicles now ferry astronauts to the ISS. Attention has turned back to the
Moon through the Artemis programme, with crewed lunar missions planned and
Mars held as the long-term goal.
"""

COOKING_MD = """\
# The Fundamentals of French Cooking

## Mother Sauces

Classical French cuisine rests on five mother sauces, codified by Auguste
Escoffier in the early twentieth century. Béchamel is milk thickened with a
white roux. Velouté is a light stock — chicken, veal, or fish — thickened with
a blond roux. Espagnole is a rich brown sauce built on roasted veal stock and
tomato. Hollandaise is an emulsion of egg yolk and clarified butter sharpened
with lemon juice. Sauce tomate is, as the name suggests, a tomato sauce
traditionally fortified with salt pork and vegetables.

## Knife Skills

A sharp chef's knife is the most important tool in the kitchen. The basic cuts
every cook should master are the julienne (matchsticks of 2 mm), the brunoise
(2 mm dice cut from julienne), the small dice (6 mm), and the chiffonade
(ribbons of leafy herbs or greens). Consistent knife cuts are not merely
cosmetic: uniform pieces cook at a uniform rate.

## Stocks and Reductions

Stock is the backbone of French sauce making. A white chicken stock simmers
raw bones gently for three to four hours with mirepoix — two parts onion, one
part carrot, one part celery — and a bouquet garni of thyme, bay leaf, and
parsley stems. Never boil a stock: agitation emulsifies fat and clouds the
liquid. A demi-glace reduces brown stock and espagnole by half until it coats
the back of a spoon.

## Technique Over Recipes

The French culinary tradition emphasises technique over recipes. A cook who
can sweat aromatics without colour, deglaze a fond with wine, mount a sauce
with cold butter, and season in layers can improvise a respectable dinner from
almost any ingredients. Mise en place — everything prepared and in its place
before the heat goes on — is the discipline that makes restaurant kitchens
run.
"""

ML_TEXT = """\
An Introduction to Machine Learning

Supervised Learning

Machine learning is the study of algorithms that improve automatically through
experience. In supervised learning, a model is trained on labelled examples:
each training instance pairs an input with the desired output. Classification
tasks predict discrete categories, such as whether an email is spam, while
regression tasks predict continuous values, such as tomorrow's temperature.
Common supervised algorithms include linear regression, logistic regression,
decision trees, random forests, gradient boosted trees, and neural networks.

Neural Networks and Deep Learning

Artificial neural networks are composed of layers of interconnected units that
apply weighted sums followed by nonlinear activation functions. Deep learning
refers to networks with many stacked layers, which learn hierarchical feature
representations directly from raw data. Convolutional neural networks excel at
image recognition by exploiting spatial locality, while transformer
architectures dominate natural language processing through self-attention,
which lets every token attend to every other token in a sequence.

Training and Generalisation

Models are trained by minimising a loss function with gradient descent. The
backpropagation algorithm computes the gradient of the loss with respect to
every weight in the network efficiently. A central challenge is
generalisation: a model that memorises its training data but fails on new
data is said to overfit. Practitioners combat overfitting with regularisation
techniques such as weight decay, dropout, early stopping, and data
augmentation, and they estimate generalisation error on held-out validation
sets.

Embeddings and Vector Search

Modern machine learning systems represent words, sentences, images, and whole
documents as dense numeric vectors called embeddings. Semantically similar
items map to nearby points in the embedding space, which makes nearest
neighbour search a powerful retrieval primitive. Vector databases index
millions of embeddings with structures such as HNSW graphs and IVF partitions
so that approximate nearest neighbour queries return in milliseconds.
"""

FINANCE_DOCX = """\
Quarterly Financial Report: Meridian Robotics Ltd

Executive Summary

Meridian Robotics closed the second quarter with revenue of 48.2 million
dollars, an increase of 14 percent year over year, driven primarily by strong
demand for the warehouse automation product line. Gross margin improved to 61
percent from 57 percent a year earlier as component costs eased and the
manufacturing team completed the transition to the consolidated assembly
facility in Austin.

Revenue Analysis

The industrial automation segment contributed 29.5 million dollars, up 18
percent, with particularly strong bookings from logistics customers in Europe.
The healthcare robotics segment grew 9 percent to 12.1 million dollars.
Service and maintenance revenue of 6.6 million dollars was flat quarter over
quarter but carries the highest margin in the portfolio at 74 percent.

Operating Expenses

Research and development spending rose to 9.8 million dollars, reflecting the
hiring of forty-two engineers for the autonomous navigation programme. Sales
and marketing expense of 7.2 million dollars held steady at 15 percent of
revenue. General and administrative costs declined 6 percent following the
completion of the enterprise resource planning migration.

Outlook

Management raises full-year revenue guidance to a range of 196 to 202 million
dollars and expects operating margin between 18 and 20 percent. Key risks to
the outlook include semiconductor lead times, foreign exchange volatility in
the euro, and the pending tariff review on imported precision actuators.
"""

CLIMATE_HTML = """\
<html>
<head><title>Understanding Ocean Currents and Climate</title></head>
<body>
<h1>Understanding Ocean Currents and Climate</h1>

<h2>The Global Conveyor Belt</h2>
<p>Ocean currents redistribute enormous quantities of heat around the planet.
The thermohaline circulation, often called the global conveyor belt, is driven
by differences in water density caused by temperature and salinity. Warm
surface water flows poleward, cools, grows denser, and sinks into the deep
ocean, where it creeps back toward the equator over centuries. A single parcel
of water can take roughly a thousand years to complete the full circuit.</p>

<h2>The Gulf Stream</h2>
<p>The Gulf Stream carries warm tropical water from the Gulf of Mexico along
the eastern coast of North America and across the Atlantic toward western
Europe. It transports about 30 million cubic metres of water per second —
more than a hundred times the flow of all the world's rivers combined. This
heat delivery is the reason winters in London are far milder than winters in
Newfoundland at the same latitude.</p>

<h2>El Niño and La Niña</h2>
<p>The El Niño Southern Oscillation is a periodic warming and cooling of the
tropical Pacific that reshapes weather worldwide. During El Niño years, weak
trade winds allow warm water to pool off South America, bringing drought to
Australia and heavy rain to Peru. La Niña reverses the pattern. These
oscillations occur every two to seven years and are the strongest single
source of year-to-year climate variability.</p>

<h2>Currents in a Warming World</h2>
<p>Climate models suggest that melting Greenland ice could freshen the North
Atlantic and weaken the Atlantic meridional overturning circulation. A
significant slowdown would cool northwestern Europe, shift tropical rain
belts, and accelerate sea-level rise along the North American coast, making
the stability of ocean circulation one of the most consequential open
questions in climate science.</p>
</body>
</html>
"""

CODE_PY = '''\
"""A small event-driven task scheduler with priority queues and retries."""

import heapq
import time
from dataclasses import dataclass, field
from typing import Callable


@dataclass(order=True)
class ScheduledTask:
    """A task queued for execution at a specific time with a priority."""

    run_at: float
    priority: int
    action: Callable[[], None] = field(compare=False)
    max_retries: int = field(default=3, compare=False)
    attempts: int = field(default=0, compare=False)


class TaskScheduler:
    """Execute callables at scheduled times, honouring priority and retries.

    Tasks are stored in a binary heap ordered by (run_at, priority). Failed
    tasks are re-queued with exponential backoff until max_retries is
    exhausted.
    """

    def __init__(self, backoff_base: float = 2.0) -> None:
        self._heap: list[ScheduledTask] = []
        self._backoff_base = backoff_base
        self.completed: list[ScheduledTask] = []
        self.failed: list[ScheduledTask] = []

    def schedule(self, action: Callable[[], None], delay: float = 0.0, priority: int = 0) -> None:
        """Queue *action* to run *delay* seconds from now."""
        task = ScheduledTask(run_at=time.monotonic() + delay, priority=priority, action=action)
        heapq.heappush(self._heap, task)

    def run_pending(self) -> int:
        """Run every task whose scheduled time has passed. Returns run count."""
        executed = 0
        now = time.monotonic()
        while self._heap and self._heap[0].run_at <= now:
            task = heapq.heappop(self._heap)
            try:
                task.action()
                self.completed.append(task)
                executed += 1
            except Exception:
                task.attempts += 1
                if task.attempts > task.max_retries:
                    self.failed.append(task)
                else:
                    task.run_at = now + self._backoff_base**task.attempts
                    heapq.heappush(self._heap, task)
        return executed
'''

INVENTORY_ROWS = [
    ("SKU", "Product", "Category", "Warehouse", "Quantity", "Unit Price"),
    ("TL-4501", "Cordless Drill 18V", "Power Tools", "Rotterdam", 240, 89.99),
    ("TL-4502", "Orbital Sander", "Power Tools", "Rotterdam", 118, 54.50),
    ("GD-2210", "Garden Hose 25m", "Garden", "Lyon", 402, 21.95),
    ("GD-2214", "Pruning Shears", "Garden", "Lyon", 655, 13.25),
    ("EL-7734", "LED Floodlight 50W", "Electrical", "Prague", 310, 32.00),
    ("EL-7801", "Extension Reel 40m", "Electrical", "Prague", 95, 61.75),
    ("PL-1120", "PVC Pipe 32mm x 2m", "Plumbing", "Rotterdam", 1280, 4.10),
    ("PL-1188", "Ball Valve Brass 1in", "Plumbing", "Prague", 233, 9.85),
]


def make_markdown_and_text() -> None:
    (FIXTURES_DIR / "space_exploration.md").write_text(SPACE_MD, encoding="utf-8")
    (FIXTURES_DIR / "french_cooking.md").write_text(COOKING_MD, encoding="utf-8")
    (FIXTURES_DIR / "climate_currents.html").write_text(CLIMATE_HTML, encoding="utf-8")
    (FIXTURES_DIR / "task_scheduler.py").write_text(CODE_PY, encoding="utf-8")


def make_pdf() -> None:
    """Render the machine-learning article into a real multi-page PDF."""
    import fitz  # PyMuPDF

    doc = fitz.open()
    rect = fitz.Rect(72, 72, 523, 770)  # 1in margins on US Letter
    paragraphs = ML_TEXT.split("\n\n")
    page = doc.new_page()
    y = rect.y0
    for para in paragraphs:
        text = " ".join(para.split())
        # Estimate height, start a new page if this paragraph won't fit.
        needed = fitz.get_text_length(text, fontsize=11) / (rect.width) * 16 + 32
        if y + needed > rect.y1:
            page = doc.new_page()
            y = rect.y0
        spent = page.insert_textbox(
            fitz.Rect(rect.x0, y, rect.x1, rect.y1), text, fontsize=11, fontname="helv", lineheight=1.4
        )
        y = rect.y1 - spent + 14  # insert_textbox returns unused height
    doc.save(FIXTURES_DIR / "machine_learning.pdf")
    doc.close()


def make_docx() -> None:
    """Render the financial report into a real DOCX with heading styles."""
    import docx

    document = docx.Document()
    lines = [block.replace("\n", " ").strip() for block in FINANCE_DOCX.split("\n\n") if block.strip()]
    document.add_heading(lines[0], level=1)
    for block in lines[1:]:
        # Single short lines are the section headings in the source text.
        if len(block) < 60:
            document.add_heading(block, level=2)
        else:
            document.add_paragraph(block)
    document.save(FIXTURES_DIR / "financial_report.docx")


def make_xlsx() -> None:
    """Create a real spreadsheet of warehouse inventory."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Inventory"
    for row in INVENTORY_ROWS:
        ws.append(row)
    wb.save(FIXTURES_DIR / "warehouse_inventory.xlsx")


def main() -> None:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    make_markdown_and_text()
    make_pdf()
    make_docx()
    make_xlsx()
    for path in sorted(FIXTURES_DIR.iterdir()):
        print(f"  wrote {path.name} ({path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
