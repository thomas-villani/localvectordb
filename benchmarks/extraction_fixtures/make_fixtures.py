"""Generate the adversarial extraction-gate fixtures.

These documents are engineered to contain the extraction failure modes that
have actually bitten: hyphen-wrapped words across line breaks, multi-line
headings (both produced 2-3 spurious sections per wrapped title before all2md
1.12 learned to merge them), running page headers/footers, nested structure,
and numbered contract clauses that look like headings.

The generated binaries are COMMITTED. Regenerating them is a fixture change:
re-run ``eval_extraction.py --update`` afterwards and review the fingerprint
diff. The e2e fixtures under ``scripts/e2e/fixtures/`` cover the well-behaved
case for every format; these cover the hostile one.

Usage:
    ./.venv/Scripts/python.exe benchmarks/extraction_fixtures/make_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

FIXTURES_DIR = Path(__file__).parent

# One line per entry: explicit line placement, so the hyphen wraps and the
# two-line title are under our control, not a text-box layouter's.
PDF_PAGE_1 = [
    ("Attention-Augmented Retrieval over Long Docu-", 16, True),
    ("ments with Compressed Section Mem-", 16, True),
    ("ory Banks", 16, True),
    ("", 11, False),
    ("A. Researcher and B. Colleague", 11, False),
    ("", 11, False),
    ("1 Introduction", 13, True),
    ("", 11, False),
    ("Retrieval systems routinely discard docu-", 11, False),
    ("ment structure. We propose a hierarchi-", 11, False),
    ("cal index whose sections carry their own vec-", 11, False),
    ("tors, and measure the effect of atten-", 11, False),
    ("tion pooling on long-span representa-", 11, False),
    ("tions across three corpora.", 11, False),
    ("", 11, False),
    ("2 Related Work", 13, True),
    ("", 11, False),
    ("Prior work on passage retrieval concen-", 11, False),
    ("trates on fixed-size windows. Sentence-", 11, False),
    ("aware chunking preserves boundaries but ig-", 11, False),
    ("nores headings entirely.", 11, False),
]

PDF_PAGE_2 = [
    ("3 Method", 13, True),
    ("", 11, False),
    ("Our encoder reads each section indepen-", 11, False),
    ("dently, then pools window vectors by weighted", 11, False),
    ("mean. The pooling temperature is corpus-", 11, False),
    ("dependent and set by validation.", 11, False),
    ("", 11, False),
    ("4 Results", 13, True),
    ("", 11, False),
    ("The hierarchical index wins on section tar-", 11, False),
    ("gets and ties on document targets, match-", 11, False),
    ("ing the intuition that structure helps precise", 11, False),
    ("queries most.", 11, False),
]

HTML_DOC = """<!DOCTYPE html>
<html><head><title>Structure Stress</title></head><body>
<h1>Structure Stress Document</h1>
<p>Exercises the constructs whose markdown rendering has changed between
all2md releases.</p>
<h2>Nested Lists</h2>
<ul>
  <li>outer one
    <ol><li>inner first</li><li>inner second</li></ol>
  </li>
  <li>outer two</li>
</ul>
<h2>A Table</h2>
<table>
  <tr><th>knob</th><th>default</th><th>regime</th></tr>
  <tr><td>chunk_size</td><td>500</td><td>encoder-dependent</td></tr>
  <tr><td>vector_weight</td><td>0.5</td><td>corpus-dependent</td></tr>
</table>
<h2>Code and Quotes</h2>
<pre><code>def fingerprint(text):
    return sha256(text)</code></pre>
<blockquote>Extraction changes move section boundaries invisibly.</blockquote>
</body></html>
"""

DOCX_CLAUSES = [
    ("Agreement of Purchase", 1),
    ("Section 1. Definitions", 2),
    (
        '"Material Adverse Effect" means any change that is materially adverse '
        "to the business, excluding changes in general economic conditions.",
        0,
    ),
    ("Section 2. Purchase and Sale", 2),
    ("Subject to Section 4.11, the Purchaser shall acquire the Shares at the " "price set out in Schedule A.", 0),
    ("Section 3. Representations", 2),
    (
        "Each party represents that the execution of this Agreement has been "
        "duly authorized, as described in Section 1 above.",
        0,
    ),
]


def make_pdf() -> None:
    import fitz  # PyMuPDF

    doc = fitz.open()
    for lines in (PDF_PAGE_1, PDF_PAGE_2):
        page = doc.new_page()
        page.insert_text((72, 40), "Preprint - under review", fontsize=8, fontname="helv")
        y = 90.0
        for text, size, bold in lines:
            if text:
                page.insert_text((72, y), text, fontsize=size, fontname="hebo" if bold else "helv")
            y += size * 1.5
        page.insert_text((280, 800), f"Page {doc.page_count}", fontsize=8, fontname="helv")
    doc.save(FIXTURES_DIR / "wrapped_headings.pdf")
    doc.close()


def make_html() -> None:
    (FIXTURES_DIR / "structure_stress.html").write_text(HTML_DOC, encoding="utf-8")


def make_docx() -> None:
    import docx

    d = docx.Document()
    for text, level in DOCX_CLAUSES:
        if level:
            d.add_heading(text, level=level)
        else:
            d.add_paragraph(text)
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Schedule"
    table.rows[0].cells[1].text = "Contents"
    table.rows[1].cells[0].text = "A"
    table.rows[1].cells[1].text = "Price per share"
    d.save(str(FIXTURES_DIR / "clause_contract.docx"))


if __name__ == "__main__":
    make_pdf()
    make_html()
    make_docx()
    print(f"fixtures written to {FIXTURES_DIR}")
